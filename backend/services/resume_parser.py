import io
import magic
from typing import Tuple, Optional

import pdfplumber
from docx import Document
import PyPDF2

from backend.utils.file_utils import(
    FileParsingError, 
    TextExtractionError, 
    FileUploadError, 
    log_error, 
    log_warning, 
    log_info, 
    with_fallback
)
from backend.core.logger import setup_logger
from backend.core.config import (
    MAX_FILE_SIZE_BYTES,
    MAX_FILE_SIZE_MB, 
    SUPPORTED_MIME_TYPES
)

logger = setup_logger("ats_resume_scorer | resume_parser")

class FileParsingError(Exception):
    """Custom exception for file parsing errors."""
    pass

class FileValidationError(Exception):
    """Custom exception for file validation errors."""
    pass

# 1. validate file size and type
def validate_file(file_data:bytes, filename:str) -> Tuple[bool, str, Optional[str]]:
    file_size_byte = len(file_data)
    
    # 1. Check for empty file
    if file_size_byte == 0:
        log_warning(f"File {filename} is empty.")
        return False, "The uploaded file is empty. Please upload a valid resume.", None
        
    # 2. Check file size limits
    if file_size_byte > MAX_FILE_SIZE_BYTES:
        size_mb = file_size_byte / (1024 * 1024)
        log_warning(f"File {filename} is too large: {size_mb:.2f} MB. Max allowed is {MAX_FILE_SIZE_MB} MB.")
        return False, f"File size ({size_mb:.2f} MB) exceeds the maximum of {MAX_FILE_SIZE_MB} MB. Please upload a smaller file or compress your resume.", None
      
    
    try:
        mime_type = magic.from_buffer(file_data, mime=True)
    except Exception as e:
        log_error(f"Error determining MIME type for file '{filename}': {str(e)}")
        return False, f'Could not determine the file type.{e}', None
    
    if mime_type not in SUPPORTED_MIME_TYPES:
        supported = ", ".join(SUPPORTED_MIME_TYPES.keys()).upper()
        return False, (
            f'Unsupported file type: {mime_type}. '
            f'Please upload one of: {supported}.'
        ), None
        
    return True, "", SUPPORTED_MIME_TYPES[mime_type]

# 2. extract hyperlinks from text
def _extract_pdf_hyperlinks(file_data:bytes)->str:
    urls = []
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_data))
        for page in reader.pages:
            annots = page.get('/Annots', [])
            if not annots:
                continue
            for annot_ref in annots:
                try:
                    annot = annot_ref.get_object()
                    # Skip if it is not a hyperlink
                    if annot.get('/Subtype') != '/Link':
                        continue
                        
                    # Safely drill down to the URL string
                    action = annot.get('/A', {})
                    uri = action.get('/URI', '')
                    
                    if uri:
                        if isinstance(uri, bytes):
                            uri = uri.decode('utf-8', errors='ignore')
                        uri = uri.strip()
                        
                        if uri.startswith(('http://', 'https://')):
                            urls.append(uri)
                except Exception as item_error:
                    # Logs individual corruption but keeps processing other links
                    log_error(f"Skipping corrupted annotation: {item_error}")
                    
    except Exception as doc_error:
        log_error(f"Failed to read PDF structure: {doc_error}")
        
    return '\n'.join(urls)

# 2. _append_hyperlinks
def _append_hyperlinks(text: str, file_data: bytes) -> str:
    """Shared helper — called once after whichever PDF extractor wins."""
    hyperlinks = _extract_pdf_hyperlinks(file_data)
    return (text.strip() + '\n' + hyperlinks) if hyperlinks else text.strip()

# 3. _extract_pdf_with_pdfplumber
def _extract_pdf_with_pdfplumber(file_data:bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(file_data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
                
        if not text.strip():
            raise TextExtractionError(
                "pdfplumber could not extract any text from the PDF. ",
                usee_message = "No text could be extracted from the PDF. "
            )
            
        return text 

# 4. _extract_pdf_with_pypdf2
def _extract_pdf_with_pypdf2(file_data: bytes) -> str:
    text = ''
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + '\n'

    if not text.strip():
        raise TextExtractionError(
            'PyPDF2 extracted no text',
            user_message='No text could be extracted from the PDF.'
        )

    return text

# 5. _extract_text_from_pdf (main entry point for pdf)
def extract_text_from_pdf(file_data: bytes) -> str:
    try:
        result, used_fallback = with_fallback(
            _extract_pdf_with_pdfplumber,
            _extract_pdf_with_pypdf2,
            file_data,
            log_fallback=True,
        )
        
        if used_fallback:
                logger.warning("pdfplumber failed to extract text, used PyPDF2 as fallback.")
                
        return _append_hyperlinks(result, file_data)
    except TextExtractionError as e:
        log_error(e, context='extract_text_from_pdf')
        raise FileParsingError(
            'Failed to extract text from PDF using both pdfplumber and PyPDF2. '
            'The PDF may be corrupted, password-protected, or contain only scanned images. '
            'Please ensure it contains selectable text.'
        ) from e

# 6. _extract_docx_hyperlinks
def _extract_docx_hyperlinks(doc: Document) -> list[str]:
    urls = []
    try:
        for rel in doc.part.rels.values():
            if 'hyperlink' in rel.reltype.lower():
                url = rel._target
                if isinstance(url, str) and url.startswith('http'):
                    urls.append(url)
    except Exception:
        pass
    return urls

# 7. _extract_text_from_docx (main entry point for docx)
def extract_text_from_docx(file_data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_data))
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        
        seen: set[str] = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct and ct not in seen:
                        seen.add(ct)
                        text_parts.append(ct)
        
        text = '\n'.join(text_parts)
        
        if not text.strip():
            raise TextExtractionError(
                'No text extracted from DOCX',
                user_message='No text could be extracted from the DOCX file.'
            )
        
        urls = _extract_docx_hyperlinks(doc)
        if urls:
            text += text.strip() + '\n' + '\n'.join(urls)
            
        log_info(f'Extracted {len(text)} chars from DOCX', context='resume_parser')
        return text.strip()
    
    except FileParsingError:
        raise
    except Exception as e:
        log_error(e, context='extract_text_from_docx')
        raise FileParsingError(
            'Failed to extract text from DOCX. '
            'The document may be corrupted or in an unsupported format. '
            'Please try re-saving or converting to PDF.'
        ) from e

# 8. _extract_text_from_doc (main entry point for doc)
def extract_text_from_doc(file_data: bytes) -> str:
    raise FileParsingError(
        'Legacy .doc format is not supported. '
        'Please convert your document to .docx or .pdf and try again. '
        'You can convert using Microsoft Word, Google Docs, or online tools.'
    ) 

# 9. extract_text
def extract_text(file_data:bytes, file_type:str)->str:
    if file_type=='pdf':
        return extract_text_from_pdf(file_data)
    elif file_type=='docx':
        return extract_text_from_docx(file_data)
    elif file_type=='doc':
        return extract_text_from_doc(file_data)
    else:
        raise FileValidationError(
            f'invalid file type: {file_type}. supported types are: pdf, docx and doc'
        )
             
# 10. parse_resume_file (main entry point for the whole process)
def parse_resume_file(file_data: bytes, filename: str)-> str:
    log_info(f"Starting to parse file: {filename}", context='resume_parser')
    
    # 1. Validate file
    try:
        is_valid, error_msg, file_type = validate_file(file_data, filename)
        if not is_valid:
            log_warning(f"File validation failed for {filename}: {error_msg}")
            raise FileValidationError(error_msg)
    except FileValidationError as ve:
        raise ve
    
    except Exception as e:
        log_error(e, context='parse_resume_file')
        raise FileUploadError(
            'An unexpected error occurred during file validation. Please try again or contact support.'
        ) from e
        
        
    # 2. Extract text
    try:
        text = extract_text(file_data, file_type)
        log_info(f"Successfully extracted text from {filename}", context='resume_parser')
    except (FileParsingError, TextExtractionError) as pe:
        raise pe
    except Exception as e:
        log_error(e, context='parse_resume_file')
        raise FileParsingError(
            'An unexpected error occurred during text extraction. Please ensure your file is not corrupted and try again.'
        ) from e
        
    metadata = {
        'filename':        filename,
        'file_type':       file_type,
        'file_size_bytes': len(file_data),
        'text_length':     len(text),
        'success':         True,
    }    
    
    return text, metadata

## 